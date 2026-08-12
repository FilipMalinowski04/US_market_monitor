import os
import tempfile

import matplotlib.dates as mdates
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

from .yfinance_service import (
    DEFAULT_PRICE_PERIOD,
    PRICE_PERIODS,
    get_company_info,
    get_price_history,
    normalize_price_period,
)


class CompanyReportGenerator:
    def __init__(self, ticker: str, period: str = DEFAULT_PRICE_PERIOD):
        self.ticker = ticker.strip().upper()
        self.period = normalize_price_period(period)
        self.period_label = PRICE_PERIODS[self.period]
        self.info = None
        self.price_data = None
        self.temp_chart_path = None
        self.temp_doc_path = None

    def fetch_data(self):
        self.info = get_company_info(self.ticker)
        if not self.info:
            raise ValueError(f"Nie znaleziono spółki: {self.ticker}")

        self.price_data = get_price_history(self.ticker, period=self.period)
        if self.price_data.empty:
            raise ValueError(f"Brak danych cenowych dla: {self.ticker}")

    def generate_chart(self):
        fig, ax = plt.subplots(figsize=(8, 4))
        ax.plot(self.price_data.index, self.price_data["Close"], color="#0d6efd")
        ax.set_title(f"{self.ticker} — cena zamknięcia ({self.period_label})")
        ax.set_ylabel("Cena (USD)")
        ax.xaxis.set_major_formatter(mdates.DateFormatter("%Y-%m"))
        fig.autofmt_xdate()
        ax.grid(True, alpha=0.3)

        fd, self.temp_chart_path = tempfile.mkstemp(suffix=".png")
        os.close(fd)
        fig.savefig(self.temp_chart_path, bbox_inches="tight", dpi=150)
        plt.close(fig)

    def create_report(self) -> str:
        doc = Document()
        name = self.info["name"]
        doc.add_heading(f"Raport spółki: {name} ({self.ticker})", level=0)

        doc.add_heading("Przegląd", level=1)
        doc.add_paragraph(
            f"Niniejszy raport zawiera podstawowe informacje o spółce {name} "
            f"oraz analizę cen akcji za okres: {self.period_label}."
        )

        doc.add_heading("Dane podstawowe", level=2)
        doc.add_paragraph(f"Symbol giełdowy: {self.ticker}")
        doc.add_paragraph(f"Nazwa: {name}")
        doc.add_paragraph(f"Kraj: {self.info['country']}")
        doc.add_paragraph(f"Sektor: {self.info['sector']}")
        doc.add_paragraph(f"Giełda: {self.info['exchange']}")
        doc.add_paragraph(f"Zakres analizy cen: {self.period_label}")

        doc.add_heading("Opis biznesowy", level=2)
        doc.add_paragraph(self.info["description"])

        latest_close = self.price_data["Close"].iloc[-1]
        earliest_close = self.price_data["Close"].iloc[0]
        change_pct = ((latest_close - earliest_close) / earliest_close) * 100

        doc.add_heading("Podsumowanie cen", level=2)
        doc.add_paragraph(
            f"Cena zamknięcia na początku okresu: ${earliest_close:.2f}. "
            f"Ostatnia cena zamknięcia: ${latest_close:.2f}. "
            f"Zmiana w okresie {self.period_label}: {change_pct:+.2f}%."
        )

        doc.add_heading(f"Wykres cen — {self.period_label}", level=2)
        doc.add_picture(self.temp_chart_path, width=Inches(6))

        fd, self.temp_doc_path = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        doc.save(self.temp_doc_path)
        return self.temp_doc_path

    def cleanup(self):
        for path in (self.temp_chart_path, self.temp_doc_path):
            if path and os.path.exists(path):
                os.remove(path)

    def generate(self) -> tuple[str, str]:
        try:
            self.fetch_data()
            self.generate_chart()
            doc_path = self.create_report()
            filename = f"raport_{self.ticker}_{self.period}.docx"
            return doc_path, filename
        except Exception:
            self.cleanup()
            raise
